import os
from docx import Document
from docx.shared import Inches
import pyautogui

# Set the folder containing images
image_folder = 'G:\Pycharm\pythonProject3\cdlpyautogui\images'  # Replace with the actual folder path
doc_filename = "Detected_Answers_org.docx"

# Create a new Word document
doc = Document()
doc.add_heading('Detected Screenshots', level=1)

# Get list of image files (sort to maintain order)
image_files = sorted([f for f in os.listdir(image_folder) if f.endswith(('.png', '.jpg', '.jpeg'))])

# Insert images into the document
for image_file in image_files:
    image_path = os.path.join(image_folder, image_file)
    doc.add_paragraph(f"Image: {image_file}")
    doc.add_picture(image_path, width=Inches(4), height=Inches(3))  # Adjust size

# Save the document
doc.save(doc_filename)
pyautogui.alert(f"Document saved as {doc_filename}", "Process Completed")

print(f"Document '{doc_filename}' created successfully with {len(image_files)} images.")
