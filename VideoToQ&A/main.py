import pyautogui
import pytesseract
import cv2
import numpy as np
import time
from PIL import ImageGrab
from docx import Document

# Configure Tesseract OCR path if needed
# Example (Windows):
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Load reference images
correct_answer_img = cv2.imread("correct_answer.png", cv2.IMREAD_GRAYSCALE)
explanation_img = cv2.imread("explanation.png", cv2.IMREAD_GRAYSCALE)
answers_img = cv2.imread("answers.png", cv2.IMREAD_GRAYSCALE)

# Create a new Word document
doc = Document()
doc.add_heading('Detected Answers', level=1)

# Variables to track the last captured question and time
last_question_text = ""
last_capture_time = 0
capture_interval = 10  # Interval in seconds between allowed captures for the same question

# Function to capture screen and check for text
def capture_screen():
    global last_question_text, last_capture_time

    screenshot = pyautogui.screenshot()
    screenshot_cv = np.array(screenshot)
    screenshot_gray = cv2.cvtColor(screenshot_cv, cv2.COLOR_BGR2GRAY)

    # OCR text detection
    text = pytesseract.image_to_string(screenshot)

    # Check for template match
    res1 = cv2.matchTemplate(screenshot_gray, correct_answer_img, cv2.TM_CCOEFF_NORMED)
    res2 = cv2.matchTemplate(screenshot_gray, explanation_img, cv2.TM_CCOEFF_NORMED)
    res3 = cv2.matchTemplate(screenshot_gray, answers_img, cv2.TM_CCOEFF_NORMED)

    threshold = 0.8
    current_time = time.time()

    # If enough time has passed and text contains the question, capture screenshot
    if (current_time - last_capture_time > capture_interval) and (text_contains_keywords(text) or np.max(res1) > threshold or np.max(res2) > threshold or np.max(res3) > threshold):
        # Check if the question is different from the last captured one
        if text != last_question_text:
            save_screenshot(screenshot)
            last_question_text = text  # Update the last question text
            last_capture_time = current_time  # Update the last capture time

# Function to check text content
def text_contains_keywords(text):
    return "Correct Answer:" in text or "Explanation" in text or "Answer(s)" in text

# Function to save screenshot to Word document
import os
from docx.shared import Inches  # Import Inches to set size in inches

# Check if the file exists, open it for appending; if not, create a new one
doc_filename = "Detected_Answers.docx"
if os.path.exists(doc_filename):
    doc = Document(doc_filename)  # Open the existing document
else:
    doc = Document()  # Create a new one

doc.add_heading('Detected Answers', level=1)  # Add title if the document is new

def save_screenshot(screenshot):
    filename = f"screenshot_{int(time.time())}.png"
    screenshot.save(filename)

    # Resize the image (adjust the width and height as needed)
    doc.add_paragraph(f"Screenshot taken at {time.strftime('%Y-%m-%d %H:%M:%S')}")

    # Add the image with desired size (width and height in inches)
    doc.add_picture(filename, width=Inches(4), height=Inches(3))  # Adjust size here

    # Save the document after adding content
    doc.save(doc_filename)
    print(f"Saved: {filename}")

# Main loop to monitor the screen
print("Monitoring for text...")
while True:
    capture_screen()
    time.sleep(2)  # Adjust interval as needed
